"""Trich xuat noi dung Final_Report_QSVM.docx ra markdown de audit.

Output:
- docs/_Final_Report_extracted.md (text noi dung + heading markdown)
- docs/_Final_Report_structure.txt (cau truc heading + so luong figure/table)
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# Duong dan tai lieu
ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "docs" / "Final_Report_QSVM.docx"
OUT_MD = ROOT / "docs" / "_Final_Report_extracted.md"
OUT_STRUCT = ROOT / "docs" / "_Final_Report_structure.txt"


def style_level(style_name: str) -> int:
    """Tra ve level heading dua tren ten style. 0 neu khong phai heading."""
    if not style_name:
        return 0
    name = style_name.lower()
    if name.startswith("heading"):
        parts = name.split()
        if len(parts) >= 2 and parts[1].isdigit():
            return int(parts[1])
        return 1
    return 0


def paragraph_has_inline_image(p_elem) -> bool:
    """Kiem tra paragraph co chua hinh anh inline thuc su (not just hyperlink/symbol)."""
    # Tim element <w:drawing> hoac <pic:pic> truc tiep, khong dung 'in xml string'
    drawings = p_elem.findall(".//" + qn("w:drawing"))
    if drawings:
        return True
    return False


def iter_block_items(doc: Document):
    """Duyet ca paragraph va table theo dung thu tu xuat hien."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            yield "p", child
        elif tag == qn("w:tbl"):
            yield "t", child


def main() -> int:
    if not DOCX_PATH.exists():
        print(f"[ERR] Khong tim thay {DOCX_PATH}")
        return 1

    doc = Document(str(DOCX_PATH))
    paragraphs = {p._element: p for p in doc.paragraphs}
    tables = {t._element: t for t in doc.tables}

    md_lines: list[str] = []
    structure_lines: list[str] = []
    fig_count = 0
    tbl_count = 0
    word_count = 0
    heading_count_by_level: dict[int, int] = {}

    for kind, el in iter_block_items(doc):
        if kind == "p":
            para = paragraphs.get(el)
            if para is None:
                continue
            text = para.text.strip()
            style_name = para.style.name if para.style else ""
            level = style_level(style_name)
            has_image = paragraph_has_inline_image(el)

            if level > 0 and text:
                heading_count_by_level[level] = heading_count_by_level.get(level, 0) + 1
                md_lines.append(f"\n{'#' * level} {text}\n")
                structure_lines.append(f"H{level:<2}      | {text[:140]}")
                continue

            if has_image:
                fig_count += 1
                md_lines.append(f"\n![FIGURE_{fig_count}](FIGURE_{fig_count})")
                structure_lines.append(f"FIGURE   #{fig_count}  | caption-near: {text[:120]!r}")
                if text:
                    md_lines.append(text)
                continue

            if not text:
                continue
            md_lines.append(text)
            word_count += len(text.split())
        elif kind == "t":
            table = tables.get(el)
            if table is None:
                continue
            tbl_count += 1
            md_lines.append(f"\n[TABLE_{tbl_count}]")
            structure_lines.append(f"TABLE    #{tbl_count}  | rows={len(table.rows)}, cols={len(table.columns)}")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
            md_lines.append("")

    OUT_MD.write_text("\n".join(md_lines), encoding="utf-8")
    OUT_STRUCT.write_text("\n".join(structure_lines), encoding="utf-8")

    print(f"[OK] Da extract: {OUT_MD}")
    print(f"[OK] Cau truc:   {OUT_STRUCT}")
    print(f"Word count (paragraph text): ~{word_count}")
    print(f"Figures: {fig_count}")
    print(f"Tables: {tbl_count}")
    print("Headings by level:")
    for lvl in sorted(heading_count_by_level.keys()):
        print(f"  H{lvl}: {heading_count_by_level[lvl]}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
