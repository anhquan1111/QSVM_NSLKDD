"""Inspect cau truc Final_Report_QSVM.docx de hieu format."""
from __future__ import annotations
from collections import Counter
from pathlib import Path
import docx

# Doc tai lieu
ROOT = Path(__file__).resolve().parent.parent
doc = docx.Document(str(ROOT / "docs" / "Final_Report_QSVM.docx"))

# Ghi het ra file de tranh van de encoding stdout
out_lines: list[str] = []
out_lines.append(f"Total paragraphs: {len(doc.paragraphs)}")
out_lines.append(f"Total tables: {len(doc.tables)}")
out_lines.append("")
out_lines.append("=== Sample of first 80 non-empty paragraphs with styles ===")

shown = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else "None"
    out_lines.append(f"[{style:25s}] {text[:160]}")
    shown += 1
    if shown >= 80:
        break

out_lines.append("")
out_lines.append("=== Style name distribution (non-empty paragraphs) ===")
counter = Counter(p.style.name for p in doc.paragraphs if p.text.strip())
for k, v in counter.most_common():
    out_lines.append(f"{k:30s} {v}")

# Phat hien heading bang text-based heuristic (so thu tu chuong, viet hoa)
out_lines.append("")
out_lines.append("=== Candidate headings by text heuristic (paragraphs <120 chars, all-caps OR starts with chapter pattern) ===")
import re
patterns = [
    re.compile(r"^(chuong|chương|chapter|phan|section|c1|c2|c3|c4|c5|c6)\b", re.IGNORECASE),
    re.compile(r"^\d+(\.\d+)*[\.\s]"),
    re.compile(r"^[IVXLC]+\.\s"),
]
count = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if not text or len(text) > 120:
        continue
    matched = any(pat.match(text) for pat in patterns) or (text.isupper() and len(text) > 4 and len(text) < 80)
    if matched:
        out_lines.append(f"  -> {text[:160]}")
        count += 1
        if count >= 100:
            break

(ROOT / "docs" / "_docx_inspect.txt").write_text("\n".join(out_lines), encoding="utf-8")
print("Wrote docs/_docx_inspect.txt")
